from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from io import BytesIO
import numpy as np
from PIL import Image

import fitz # PyMuPDF
import pandas as pd
from docx import Document
from docx.shared import Pt
import tempfile
import os
from fastapi.responses import FileResponse


app = FastAPI()

# Allow CORS for local dev
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Try to load EasyOCR — it may not be installed yet
reader = None
try:
    import easyocr
    print("Loading EasyOCR model...")
    reader = easyocr.Reader(['en'], gpu=False)
    print("EasyOCR model loaded successfully.")
except ImportError:
    print("⚠️  EasyOCR not installed yet. OCR endpoint will return an error until it's installed.")
except Exception as e:
    print(f"⚠️  EasyOCR failed to load: {e}")

@app.get("/api/health")
async def health():
    return {"status": "ok", "ocr_ready": reader is not None}

@app.post("/api/ocr")
async def process_ocr(file: UploadFile = File(...)):
    if reader is None:
        return {"success": False, "error": "EasyOCR is not installed yet. Run: pip install easyocr"}
    try:
        contents = await file.read()
        image = Image.open(BytesIO(contents))
        img_np = np.array(image)
        
        results = reader.readtext(img_np)
        extracted_text = "\n".join([text for (_, text, _) in results])
        
        return {"success": True, "text": extracted_text}
    except Exception as e:
        return {"success": False, "error": str(e)}



@app.post("/api/ocr/analyze")
async def analyze_ocr(file: UploadFile = File(...)):
    if reader is None:
        return {"success": False, "error": "EasyOCR is not installed yet."}
    try:
        contents = await file.read()
        image = Image.open(BytesIO(contents))
        # Ensure image is in RGB
        if image.mode != 'RGB':
            image = image.convert('RGB')
            
        img_np = np.array(image)
        
        # We use paragraph=False to ensure granular bounding boxes.
        results = reader.readtext(img_np, paragraph=False)
        
        blocks = []
        for (bbox, text, prob) in results:
            # bbox is [[x1, y1], [x2, y1], [x2, y2], [x1, y2]]
            x1, y1 = bbox[0]
            x2, y2 = bbox[2]
            
            # Convert numpy types to native Python types for JSON serialization
            blocks.append({
                "text": text,
                "x": float(x1),
                "y": float(y1),
                "width": float(x2 - x1),
                "height": float(y2 - y1),
                "prob": float(prob)
            })
            
        return {"success": True, "blocks": blocks}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/convert/xlsx")
async def convert_to_xlsx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        doc = fitz.open(stream=contents, filetype="pdf")
        
        all_tables = []
        for page in doc:
            tabs = page.find_tables()
            if tabs:
                for tab in tabs:
                    all_tables.append(tab.extract())
        
        if not all_tables:
            return {"success": False, "error": "No tables found in this PDF"}
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp_name = tmp.name
            
        with pd.ExcelWriter(tmp_name) as writer:
            for idx, table in enumerate(all_tables):
                # Filter out None and empty rows
                clean_table = [[str(cell) if cell is not None else "" for cell in row] for row in table]
                df = pd.DataFrame(clean_table)
                df.to_excel(writer, sheet_name=f"Table_{idx+1}", index=False, header=False)
                
        return FileResponse(tmp_name, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", filename="Converted.xlsx")
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/convert/docx")
async def convert_to_docx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        doc = fitz.open(stream=contents, filetype="pdf")
        
        docx_file = Document()
        
        # Simple heuristic: extract text blocks with font sizes
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if b['type'] == 0: # Text block
                    for line in b["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if not text:
                                continue
                            
                            size = span["size"]
                            font = span["font"]
                            is_bold = "Bold" in font or "Black" in font
                            
                            if size > 16 or is_bold:
                                p = docx_file.add_heading(text, level=1 if size > 20 else 2)
                            else:
                                p = docx_file.add_paragraph(text)
                            
            docx_file.add_page_break()
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_name = tmp.name
        
        docx_file.save(tmp_name)
        
        return FileResponse(tmp_name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="Converted.docx")
    except Exception as e:
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
