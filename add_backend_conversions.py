import re

with open("backend/main.py", "r") as f:
    content = f.read()

# Add imports
imports = """
import fitz # PyMuPDF
import pandas as pd
from docx import Document
from docx.shared import Pt
import tempfile
import os
from fastapi.responses import FileResponse
"""
content = content.replace("from PIL import Image", "from PIL import Image\n" + imports)

# Add endpoints
endpoints = """
@app.post("/api/convert/xlsx")
async def convert_to_xlsx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        doc = fitz.open(stream=contents, filetype="pdf")
        
        all_tables = []
        for page in doc:
            tabs = page.find_tables()
            if tabs:
                for tab in tabs:
                    all_tables.append(tab.extract())
        
        if not all_tables:
            return {"success": False, "error": "No tables found in this PDF"}
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp_name = tmp.name
            
        with pd.ExcelWriter(tmp_name) as writer:
            for idx, table in enumerate(all_tables):
                # Filter out None and empty rows
                clean_table = [[str(cell) if cell is not None else "" for cell in row] for row in table]
                df = pd.DataFrame(clean_table)
                df.to_excel(writer, sheet_name=f"Table_{idx+1}", index=False, header=False)
                
        return FileResponse(tmp_name, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", filename="Converted.xlsx")
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/api/convert/docx")
async def convert_to_docx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        doc = fitz.open(stream=contents, filetype="pdf")
        
        docx_file = Document()
        
        # Simple heuristic: extract text blocks with font sizes
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if b['type'] == 0: # Text block
                    for line in b["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if not text:
                                continue
                            
                            size = span["size"]
                            font = span["font"]
                            is_bold = "Bold" in font or "Black" in font
                            
                            if size > 16 or is_bold:
                                p = docx_file.add_heading(text, level=1 if size > 20 else 2)
                            else:
                                p = docx_file.add_paragraph(text)
                            
            docx_file.add_page_break()
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_name = tmp.name
        
        docx_file.save(tmp_name)
        
        return FileResponse(tmp_name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="Converted.docx")
    except Exception as e:
        return {"success": False, "error": str(e)}
"""

content = content.replace("if __name__ == \"__main__\":", endpoints + "\nif __name__ == \"__main__\":")

with open("backend/main.py", "w") as f:
    f.write(content)
